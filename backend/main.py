from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, Response, FileResponse
from pydantic import BaseModel
from docx import Document
from io import BytesIO
import os
from dotenv import load_dotenv
from openai import OpenAI, AuthenticationError
import asyncio
load_dotenv()

HF_TOKEN = os.getenv("HF_TOKEN",'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')
MODEL = os.getenv("MODEL", "openai/gpt-oss-120b:cerebras")
client = OpenAI(
    base_url="https://router.huggingface.co/v1",
    api_key=HF_TOKEN,
)
import pathlib

app = FastAPI()
frontend_path = pathlib.Path(__file__).parent.parent / "frontend"

@app.get("/", response_class=HTMLResponse)
async def serve_index():
    index_path = frontend_path / "index.html"
    if index_path.exists():
        return index_path.read_text(encoding="utf-8")
    return HTMLResponse(content="Index file not found", status_code=404)

@app.get("/img/{filename}")
async def serve_image(filename: str):
    file_path = frontend_path / "img" / filename
    if file_path.exists():
        return FileResponse(file_path)
    return {"error": "Image not found"}


class GenerateRequest(BaseModel):
    mode: str 
    country: str
    case_description: str
    doc_type: str | None = None 
async def call_hf_chat(messages, max_tokens=512):

    def generate():
        try:
            completion = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                max_tokens=max_tokens,
            )
            return completion.choices[0].message.content
        except AuthenticationError as e:
            raise HTTPException(status_code=401, detail="Authentication failed. Please check your HF_TOKEN in the .env file.")

    result = await asyncio.to_thread(generate)
    return result


@app.post('/api/generate')
async def generate(req: GenerateRequest):
   
    if req.mode not in ('advice', 'document'):
        raise HTTPException(status_code=400, detail='mode must be "advice" or "document"')

    system_prompt = (
        "You are Legal.AI, a helpful assistant. Always be clear this is informational only and not a substitute for professional legal advice. "
        "Answer concisely and base your answer on the laws and common practice in the user's specified country unless told otherwise."
    )

    if req.mode == 'advice':
        user_content = (
            f"Country: {req.country}\nCase: {req.case_description}\n"
            "Provide a concise, step-by-step explanation of possible legal options, likely next steps, and references to typical laws/acts if known. Use plain language."
        )
        messages = [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_content}
        ]
        content = await call_hf_chat(messages, max_tokens=700)
        return JSONResponse({'type': 'advice', 'text': content})


    if req.mode == 'document':
        if not req.doc_type:
            raise HTTPException(status_code=400, detail='doc_type is required for document mode')
            
        doc_prompt = (
            f"Country: {req.country}\nDocument type: {req.doc_type}\nFacts: {req.case_description}\n"
            "Draft a formal legal document suitable for the specified country. Include placeholders for names/dates/addresses. Keep tone formal and add a short sign-off block. Output only the document body in plain text — do not wrap in quotes."
        )
        messages = [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': doc_prompt}]
        content = await call_hf_chat(messages, max_tokens=900)

        doc = Document()
        doc.add_heading(req.doc_type, level=1)
        for para in content.split('\n'):
            if para.strip(): 
                doc.add_paragraph(para)
        
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        safe_filename = req.doc_type.replace(' ', '_').replace('/', '_')
        
        return Response(
            content=bio.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f"attachment; filename={safe_filename}.docx"}
        )

@app.get('/api/health')
async def health():
    """A simple health check endpoint."""
    return {'ok': True}
@app.get("/{path:path}", response_class=HTMLResponse)
async def serve_spa(path: str):

    if path.startswith("api/"):
        return HTMLResponse(content="Not Found", status_code=404)
    index_path = frontend_path / "index.html"
    if index_path.exists():
        return index_path.read_text(encoding="utf-8")
    return HTMLResponse(content="Index file not found", status_code=404)

if __name__ == '__main__':
    import uvicorn
    uvicorn.run(app)

